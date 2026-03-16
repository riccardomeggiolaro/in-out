from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# --- Title ---
title = doc.add_heading('Sistema di Controllo Sbarre con Priorità Veicoli', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

doc.add_paragraph(
    'Sistema di gestione accessi e pesatura con controllo sbarre automatico, '
    'integrato con un sistema esterno di identificazione targhe. '
    'La priorità determina l\'ordine di servizio in coda quando più prenotazioni hanno la stessa data.'
)

# --- Fase 1 ---
doc.add_heading('1. Il Processo', level=1)
doc.add_heading('Fase 1: Arrivo del Mezzo', level=2)

items = [
    'Il veicolo arriva all\'ingresso del sito.',
    'Il sistema esterno (del cliente) identifica la targa tramite telecamera.',
    'Il sistema esterno chiama la nostra API per creare un accesso, passando:',
]
for item in items:
    doc.add_paragraph(item, style='List Number')

sub_items = [
    'Targa del veicolo',
    'Data/Ora dell\'arrivo',
    'Slot assegnato (fascia oraria o posizione)',
    'Priorità (usata per ordinare la coda in caso di stessa data)',
]
for s in sub_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(2.5)
    p.add_run(s)

p = doc.add_paragraph('L\'accesso può essere di due tipi:')
p.paragraph_format.space_before = Pt(6)

types = [
    ('Con pesatura', 'accesso normale, il veicolo verrà pesato sulla Pesa 2.'),
    ('Solo transito', 'senza pesatura, viene registrato solo il passaggio.'),
]
for bold_text, desc in types:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_text + ': ')
    run.bold = True
    p.add_run(desc)

# --- Fase 2 ---
doc.add_heading('Fase 2: Gestione Coda e Apertura Sbarra', level=2)

doc.add_paragraph(
    'Il software monitora continuamente la Pesa 1. '
    'Quando la pesa è vuota (peso sotto la soglia minima):'
)

steps = [
    'Cerca il prossimo accesso da servire:\n'
    '    • Prima passa la prenotazione con data meno recente (la più vecchia ha precedenza)\n'
    '    • A parità di data, passa quella con priorità più alta',
    'Apre la sbarra corrispondente.',
    'Mostra la targa sul pannello e attiva la sirena.',
    'Quando il veicolo sale sulla pesa (peso rilevato) → chiude la sbarra.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# --- Fase 3 ---
doc.add_heading('Fase 3: Pesatura Automatica sulla Pesa 2', level=2)

doc.add_paragraph('Per i veicoli con accesso con pesatura:')

steps = [
    'Il veicolo procede dalla Pesa 1 alla Pesa 2.',
    'Viene identificato automaticamente tramite telecamera (targa) o badge RFID.',
    'Il sistema esegue la pesata automatica: attende stabilità del peso, registra il peso, genera report.',
    'L\'accesso viene chiuso.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# --- Fase 4 ---
doc.add_heading('Fase 4: Transito Senza Pesata', level=2)

doc.add_paragraph('Per i veicoli registrati come solo transito:')

steps = [
    'La sbarra si apre normalmente seguendo l\'ordine di coda.',
    'Il veicolo transita sulla Pesa 1 ma la pesatura NON scatta.',
    'Viene registrato solo il passaggio (data/ora).',
    'Utile per veicoli di servizio o mezzi che non necessitano di pesatura.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# --- Ordinamento Coda ---
doc.add_heading('2. Ordinamento della Coda', level=1)

doc.add_paragraph('La coda viene ordinata così:')

p = doc.add_paragraph(style='List Number')
run = p.add_run('Data prenotazione ASC')
run.bold = True
p.add_run(' — la prenotazione più vecchia passa prima')

p = doc.add_paragraph(style='List Number')
run = p.add_run('Priorità DESC')
run.bold = True
p.add_run(' — a parità di data, passa prima quella con priorità più alta')

doc.add_paragraph('')  # spacer

# Tabella esempio
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Targa', 'Data Prenotazione', 'Priorità', 'Ordine di Servizio']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['AB123', '10/03 ore 08:00', '1', '1° (data più vecchia)'],
    ['CD456', '11/03 ore 09:00', '3', '2° (stessa data, priorità alta)'],
    ['EF789', '11/03 ore 09:00', '1', '3° (stessa data, priorità bassa)'],
    ['GH012', '12/03 ore 10:00', '2', '4° (data più recente)'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

# --- Schema del Processo ---
doc.add_heading('3. Schema del Processo', level=1)

schema = (
    'Sistema Esterno          Nostro Software            Pesa 1           Sbarra           Pesa 2\n'
    '     │                        │                       │                 │               │\n'
    '     │  POST /access          │                       │                 │               │\n'
    '     │  (targa, slot,         │                       │                 │               │\n'
    '     │   priorità, tipo)      │                       │                 │               │\n'
    '     │───────────────────────►│                       │                 │               │\n'
    '     │                        │  Salva accesso        │                 │               │\n'
    '     │                        │                       │                 │               │\n'
    '     │                        │  [Monitoraggio Pesa]  │                 │               │\n'
    '     │                        │◄──────────────────────│ pesa vuota      │               │\n'
    '     │                        │                       │                 │               │\n'
    '     │                        │  Cerca prossimo       │                 │               │\n'
    '     │                        │  (data ASC, poi       │                 │               │\n'
    '     │                        │   priorità DESC)      │                 │               │\n'
    '     │                        │                       │                 │               │\n'
    '     │                        │  Apri sbarra ─────────────────────────►│               │\n'
    '     │                        │                       │                 │ APERTA        │\n'
    '     │                        │  Pannello + Sirena    │                 │               │\n'
    '     │                        │                       │                 │               │\n'
    '     │                  [Veicolo sale sulla pesa]      │                 │               │\n'
    '     │                        │◄──────────────────────│ peso rilevato   │               │\n'
    '     │                        │                       │                 │               │\n'
    '     │                        │  Chiudi sbarra ───────────────────────►│ CHIUSA        │\n'
    '     │                        │                       │                 │               │\n'
    '     │                        │  ┌─ SE solo transito: │                 │               │\n'
    '     │                        │  │  Registra passaggio│                 │               │\n'
    '     │                        │  │  FINE              │                 │               │\n'
    '     │                        │  │                    │                 │               │\n'
    '     │                        │  └─ SE con pesatura:  │                 │               │\n'
    '     │                        │     Veicolo va a Pesa 2                │               │\n'
    '     │                        │                       │                 │               │\n'
    '     │                        │  Identificazione ─────────────────────────────────────►│\n'
    '     │                        │  (telecamera/badge)   │                 │               │\n'
    '     │                        │◄──────────────────────────────────────────────────────│\n'
    '     │                        │  Pesatura automatica  │                 │               │\n'
    '     │                        │  Accesso chiuso       │                 │               │\n'
)

p = doc.add_paragraph()
run = p.add_run(schema)
run.font.name = 'Consolas'
run.font.size = Pt(8)

# --- Cosa Serve ---
doc.add_heading('4. Cosa Serve Implementare', level=1)

doc.add_heading('Tabella transit_access', level=2)
doc.add_paragraph(
    'Tabella separata per i soli transiti (senza pesatura). '
    'Contiene: targa, slot, priorità, data ingresso/uscita, note, ID dal sistema esterno.'
)

doc.add_heading('API per il Sistema Esterno', level=2)
doc.add_paragraph('Endpoint per il sistema esterno per caricare gli accessi:')
apis = [
    ('POST /api/open-to-customer/access', 'Crea accesso con pesatura (con targa, slot, priorità)'),
    ('POST /api/open-to-customer/transit-access', 'Crea accesso solo transito (con targa, slot, priorità)'),
]
for endpoint, desc in apis:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(endpoint)
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(f' — {desc}')

doc.add_heading('Automatismo Apertura Sbarre', level=2)
steps = [
    'Nel Callback Realtime, quando la pesa è vuota: cercare il prossimo accesso/transito in coda (data ASC, priorità DESC).',
    'Aprire il relè della sbarra.',
    'Quando il veicolo sale sulla pesa, chiudere la sbarra.',
    'Se è un transito → registrare solo il passaggio, non avviare pesatura.',
    'Se è un accesso normale → procedere con il flusso di pesatura esistente.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# --- Esempi ---
doc.add_heading('5. Esempi Pratici', level=1)

doc.add_heading('Veicolo con pesatura', level=2)
doc.add_paragraph(
    'Sistema esterno invia targa + data + priorità → pesa vuota → apre sbarra → '
    'veicolo transita su Pesa 1 → chiude sbarra → va a Pesa 2 → pesatura automatica → accesso chiuso.'
)

doc.add_heading('Veicolo solo transito', level=2)
doc.add_paragraph(
    'Sistema esterno invia targa + data + priorità + tipo transito → pesa vuota → '
    'apre sbarra → veicolo passa → chiude sbarra → registra passaggio → nessuna pesatura.'
)

doc.add_heading('Coda con stessa data', level=2)
doc.add_paragraph(
    'Due veicoli prenotati per la stessa data, uno con priorità 3 e uno con priorità 1 → '
    'passa prima quello con priorità 3. Se hanno date diverse, passa prima quello con data '
    'più vecchia indipendentemente dalla priorità.'
)

# --- Save ---
doc.save('/home/user/in-out/docs/VEHICLE_PRIORITY_GATE_CONTROL.docx')
print("Done")
